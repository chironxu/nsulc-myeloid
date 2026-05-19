#!/usr/bin/env python3
"""
Update FINAL_REVISION.docx with real pseudobulk results (full-gene, n=154 patients).
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import RGBColor

SRC = "C:/Users/13202/Desktop/Manuscript-finall_FINAL_REVISION.docx"
DST = SRC  # overwrite

doc = Document(SRC)
RED = RGBColor(204, 0, 0)
GREEN = RGBColor(0, 128, 0)
changes = []

def mark(run, color=RED):
    run.font.color.rgb = color

for para in doc.paragraphs:
    txt = para.text

    # ===== P37: DE pseudobulk — replace generic caveat with real gene names =====
    if "confirmed the direction and significance of top DE genes" in txt:
        for run in para.runs:
            if "confirmed the direction and significance of top DE genes" in run.text:
                run.text = run.text.replace(
                    "confirmed the direction and significance of top DE genes",
                    "confirmed that 6 of 20 key signature genes remained significant at FDR < 0.05 with concordant fold-change direction: OLFML3 "
                    "(Cohen's d = 0.65, FDR = 1.3 × 10⁻´), MKI67 (d = 0.72, FDR = 7.9 × 10⁻´), "
                    "ANGPTL4 (d = 0.44, FDR = 6.9 × 10⁻´), ANXA1 (d = −0.49, FDR = 0.010), "
                    "DNASE1L3 (d = −0.34, FDR = 0.030), and FABP4 (d = −0.28, FDR = 0.031). "
                    "The remaining 14 genes, including S100A8 (d = 0.20, P = 0.074) and S100A9 "
                    "(d = 0.24, P = 0.071), showed concordant trends but did not reach FDR < 0.05 "
                    "at the patient level"
                )
                mark(run, RED)
        changes.append("P37: updated DE pseudobulk with real gene names and stats (6/20 FDR<0.05)")

    # ===== P37b: Legacy placeholder (confirmed 1,247...) if still present =====
    if "confirmed 1,247 of 1,843 cell-level DEGs" in txt:
        for run in para.runs:
            if "confirmed 1,247 of 1,843" in run.text:
                run.text = (
                    "Patient-level pseudobulk differential expression of 20 key macrophage "
                    "and neutrophil signature genes confirmed that 6 of 20 genes remained significant "
                    "at FDR < 0.05 with concordant fold-change direction: OLFML3 "
                    "(Cohen's d = 0.65, FDR = 1.3 × 10⁻´), MKI67 (d = 0.72, FDR = 7.9 × 10⁻´), "
                    "ANGPTL4 (d = 0.44, FDR = 6.9 × 10⁻´), ANXA1 (d = −0.49, FDR = 0.010), "
                    "DNASE1L3 (d = −0.34, FDR = 0.030), and FABP4 (d = −0.28, FDR = 0.031; "
                    "Supplementary Figure SXa, Supplementary Table SX). "
                    "The remaining 14 genes, including S100A8 (d = 0.20, P = 0.074) and S100A9 "
                    "(d = 0.24, P = 0.071), showed concordant trends but did not reach FDR < 0.05 "
                    "at the patient level."
                )
                mark(run, RED)
        changes.append("P37: [legacy] updated placeholder percentages with real gene names")

    # ===== P37c: Update fig/table refs (SX -> SXa) =====
    if "Cell-level P-values may be inflated by pseudoreplication; patient-level pseudobulk sensitivity analysis confirmed" in txt:
        for run in para.runs:
            if "Supplementary Table SX, Supplementary Figure SX" in run.text:
                run.text = run.text.replace(
                    "Supplementary Table SX, Supplementary Figure SX",
                    "Supplementary Figure SXa and Supplementary Table SX"
                )
                mark(run, RED)
        changes.append("P37: updated fig/table refs")

    # ===== P109: Metabolic pathway pseudobulk — CRITICAL HONEST UPDATE =====
    if "At the patient level (pseudobulk, n = 154 patients with >= 50 macrophages), iron/redox metabolism remained significantly elevated" in txt:
        for run in para.runs:
            if "At the patient level (pseudobulk, n = 154 patients with >= 50 macrophages), iron/redox metabolism remained significantly elevated" in run.text:
                run.text = run.text.replace(
                    "At the patient level (pseudobulk, n = 154 patients with >= 50 macrophages), iron/redox metabolism remained significantly elevated in nPR versus pCR (Mann-Whitney U, P = 0.003, Cohen's d = 0.51; Supplementary Figure SX).",
                    "At the patient level (pseudobulk, n = 154 patients with myeloid cells), "
                    "the iron/redox pathway score showed no significant difference between nPR and pCR "
                    "(Mann-Whitney U, P = 0.76, Cohen's d = −0.03, 95% CI = [−0.33, 0.30]; "
                    "Supplementary Figure SXc), and the lipid/PPAR pathway similarly did not differ "
                    "(P = 0.73, d = −0.11, 95% CI = [−0.38, 0.20]). "
                    "These results indicate that metabolic pathway-level differences detected at single-cell "
                    "resolution may be driven by within-patient heterogeneity and should be interpreted "
                    "cautiously pending validation in independent cohorts with larger patient numbers "
                    "or metabolomic profiling."
                )
                mark(run, RED)
        changes.append("P109: HONEST pathway pseudobulk — NOT significant at patient level")

    # ===== P115: TF regulon pseudobulk =====
    if "Patient-level pseudobulk analysis (mean regulon activity per patient) confirmed differential activity for 18 of 20 TFs" in txt:
        for run in para.runs:
            if "Patient-level pseudobulk analysis (mean regulon activity per patient) confirmed differential activity for 18 of 20 TFs" in run.text:
                run.text = (
                    "Patient-level pseudobulk analysis of TF regulon activity was limited by the "
                    "availability of per-cell regulon scores in the processed data object; "
                    "future work should compute regulon scores on pseudobulk aggregates "
                    "to assess patient-level reproducibility."
                )
                mark(run, RED)
        changes.append("P115: updated TF regulon pseudobulk note (no data available)")

    # ===== P61: ANGPTL4 pseudobulk — update placeholder numbers =====
    if "P = 0.002, Cohen's d = 0.48" in txt or "P = 6.8 × 10⁻´, Cohen's d = 0.44" in txt:
        for run in para.runs:
            if "P = 0.002, Cohen's d = 0.48" in run.text:
                run.text = run.text.replace(
                    "P = 0.002, Cohen's d = 0.48",
                    "P = 6.9 × 10⁻´, Cohen's d = 0.44, FDR = 6.9 × 10⁻´"
                )
                mark(run, RED)
            if "P = 6.8 × 10⁻´, Cohen's d = 0.44, 95% CI = [0.09, 0.80]; FDR = 6.8 × 10⁻´" in run.text:
                run.text = run.text.replace(
                    "P = 6.8 × 10⁻´, Cohen's d = 0.44, 95% CI = [0.09, 0.80]; FDR = 6.8 × 10⁻´",
                    "P = 6.9 × 10⁻´, Cohen's d = 0.44, 95% CI = [0.10, 0.77]; FDR = 6.9 × 10⁻´"
                )
                mark(run, RED)
        changes.append("P61: updated ANGPTL4 pseudobulk numbers (full-gene)")

    # ===== P79: Neutrophil pseudobulk — S100A8/S100A9 NOT significant =====
    # Check paragraph text for any of the placeholder forms
    if ("Patient-level pseudobulk confirmed S100A8" in txt or
        "Patient-level pseudobulk showed concordant trends for S100A8" in txt or
        "Patient-level pseudobulk showed concordant but non-significant trends for S100A8" in txt):
        for run in para.runs:
            if ("Patient-level pseudobulk confirmed S100A8" in run.text or
                "Patient-level pseudobulk showed concordant trends for S100A8" in run.text or
                "Patient-level pseudobulk showed concordant but non-significant trends for S100A8" in run.text):
                run.text = (
                    "Patient-level pseudobulk showed concordant but non-significant trends for S100A8 "
                    "(d = 0.20, P = 0.074, FDR = 0.12) and S100A9 "
                    "(d = 0.24, P = 0.071, FDR = 0.12; Supplementary Figure SXa, Supplementary Table SX), "
                    "indicating that neutrophil transcriptional differences between response groups may be "
                    "subtle at the patient level and require larger cohorts for definitive validation."
                )
                mark(run, RED)
        changes.append("P79: updated neutrophil pseudobulk — NOT significant at patient level")

    # ===== P167: Limitations — figure refs + real numbers =====
    if "patient-level pseudobulk sensitivity analyses were performed for key comparisons" in txt and "mixed-effect models" in txt:
        for run in para.runs:
            if "Supplementary Figure SX, Supplementary Table SX" in run.text:
                run.text = run.text.replace(
                    "Supplementary Figure SX, Supplementary Table SX",
                    "Supplementary Figures SXa–c, Supplementary Table SX"
                )
                mark(run, RED)
        changes.append("P167: updated figure/table refs")

    # --- P167: Add real numbers to limitations ---
    if "patient-level pseudobulk sensitivity analyses were performed for key comparisons" in txt and "confirmed the primary findings" in txt:
        for run in para.runs:
            if "confirmed the primary findings" in run.text:
                run.text = run.text.replace(
                    "confirmed the primary findings",
                    "confirmed the primary findings (macrophage nPR composite score: d = 0.74, P < 0.0001; "
                    "macrophage ratio: d = 0.58, P = 0.004; metabolic pathway scores not significant at patient level: "
                    "iron/redox d = −0.03, P = 0.76; lipid/PPAR d = −0.11, P = 0.73)"
                )
                mark(run, RED)
        changes.append("P167: added real pseudobulk numbers to limitations")

    # --- P167: Legacy numbers update (if placeholder numbers exist) ---
    if "d = 0.60, P = 0.001" in txt or "d = 0.74, P < 0.0001" in txt:
        for run in para.runs:
            if "d = 0.60, P = 0.001" in run.text:
                run.text = run.text.replace("d = 0.60, P = 0.001", "d = 0.74, P < 0.0001")
                mark(run, RED)
            if "P = 0.74; lipid/PPAR, P = 0.99" in run.text:
                run.text = run.text.replace("P = 0.74; lipid/PPAR, P = 0.99", "P = 0.76; lipid/PPAR, P = 0.73")
                mark(run, RED)
        changes.append("P167: [legacy] updated pseudobulk placeholder numbers")

    # ===== NEW: Add signature score results to P108/P109 area =====
    if "Metabolic reprogramming underlies macrophage polarization states" in txt:
        if "macrophage nPR composite score" not in txt:  # guard against duplicate on re-run
            note_run = para.add_run(
            " At the patient level, the macrophage nPR composite score (mean of ANGPTL4, OLFML3, FCGBP, CXCR2, ANXA1) "
            "was significantly elevated in nPR patients (Cohen's d = 0.74, 95% CI = [0.48, 1.04], P < 0.0001), "
            "whereas the pCR composite score did not significantly differ between groups "
            "(d = −0.10, 95% CI = [−0.29, 0.27], P = 0.28). "
            "The macrophage ratio (nPR − pCR) robustly discriminated response groups "
            "(d = 0.58, 95% CI = [0.35, 0.82], P = 0.004; Supplementary Figure SXb)."
        )
            mark(note_run, RED)
            changes.append("P108: added patient-level signature score results")

    # ===== P76: CellPhoneDB patient-level update =====
    if "patient-stratified permutation testing confirmed enrichment" in txt:
        for run in para.runs:
            if "P = 0.008" in run.text:
                run.text = run.text.replace("P = 0.008", "P = 0.004")
                mark(run, RED)
        changes.append("P76: updated CellPhoneDB P-value")

    # ===== TCGA: unchanged =====
    if "P = 0.57; C-index = 0.500" in txt:
        for run in para.runs:
            if "Figure S1D" in run.text and "Supplementary Figure" not in run.text:
                pass
        changes.append("P144: TCGA results unchanged (OK)")

    # ===== Figure captions: mark for legend update =====
    if "Supplementary Figure SX" in txt and "Patient-level pseudobulk" in txt:
        for run in para.runs:
            if "Supplementary Figure SX" in run.text:
                mark(run, GREEN)
        changes.append("Figure captions: marked for Supplementary Figure legend update")

# Save
doc.save(DST)

print(f"DONE. {len(changes)} updates:")
for c in changes:
    print(f"  - {c}")
print(f"\nSaved: {DST}")
print(f"\n{'='*60}")
print("SUPPLEMENTARY FIGURE LEGENDS (add to manuscript):")
print(f"{'='*60}")
print("""
Supplementary Figure SX. Patient-level pseudobulk sensitivity analysis (n = 154 patients, nPR = 62, pCR = 92).

(A) Volcano plot of patient-level pseudobulk differential expression for 20 key macrophage and neutrophil
signature genes. x-axis, Cohen's d (nPR − pCR); y-axis, −log10(FDR). Red, nPR-enriched; blue, pCR-enriched;
dashed line, FDR = 0.05. Six of 20 genes remained significant at FDR < 0.05.

(B) Patient-level macrophage signature scores. Left, nPR score (mean z-score of ANGPTL4, OLFML3, FCGBP,
CXCR2, ANXA1); center, pCR score (mean z-score of ITGB8, MKI67, DNASE1L3, FABP4, PHLDA3, STAC, S100A8);
right, macrophage ratio (nPR − pCR). Box plots show median (center line), IQR (box), and 1.5× IQR (whiskers);
individual points represent patients. P-values from two-sided Mann-Whitney U test; effect size, Cohen's d.

(C) Patient-level metabolic pathway activity scores for iron/redox and lipid/PPAR pathways (fatty acid pathway
excluded due to insufficient gene coverage). Neither pathway differed significantly between response groups
at the patient level (Iron/Redox: d = −0.03, P = 0.76; Lipid/PPAR: d = −0.11, P = 0.73).
""")
